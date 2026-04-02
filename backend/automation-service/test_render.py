from pathlib import Path
import json
import sys
import re
import os

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH  # ✅ Added JUSTIFY import

# --------------------------
# Paths
# --------------------------
TEMPLATE_PATH = Path("templates/template.docx")
JSON_PATH = Path("data/base_content.json")

# Base folder where numbered subfolders will be created
BASE_FOLDER = Path("outputs")

# Resume final filename
FINAL_FILENAME = "Ajay_Purshotam_Thota.docx"

# Temporary debug file
OUTPUT_RAW = Path("Tailored_Resume_RAW.docx")  # saved right after docxtpl render

# --------------------------
# Helpers
# --------------------------
LIST_STYLE_CANDIDATES = {
    "List Bullet",
    "List Paragraph",
    "Bullet",
    "List Bullet 2",
    "List Bullet 3",
    "Body Text List",
}

BULLET_PREFIXES = ("•", "-", "–", "—", "*")  # what we’ll auto-convert to a real bullet style


def tighten_list_paragraph(p):
    """Tighten line spacing and justify bullet paragraphs."""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  


def is_list_style(paragraph):
    """Check if a paragraph uses a list-like style."""
    try:
        return paragraph.style and paragraph.style.name in LIST_STYLE_CANDIDATES
    except Exception:
        return False


def remove_paragraph(paragraph):
    """Remove an empty or unwanted paragraph."""
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)
    p._p = p._element = None


def normalize_bullets(doc: Document):
    """Normalize bullets and ensure consistent list formatting."""
    prev_was_list = False
    to_delete = []

    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()

        # Convert manually typed bullets into Word list style
        if txt.startswith(BULLET_PREFIXES):
            raw = p.text
            new_text = raw.lstrip("".join(BULLET_PREFIXES)).lstrip(" \t")
            p.text = new_text if new_text else ""
            try:
                p.style = doc.styles["List Bullet"]
            except KeyError:
                p.style = doc.styles["List Paragraph"]

        # Tighten and justify list paragraphs
        if is_list_style(p):
            tighten_list_paragraph(p)
            if prev_was_list and txt == "":
                to_delete.append(p)
            prev_was_list = True
        else:
            prev_was_list = False

    # Remove empty bullet paragraphs
    for p in to_delete:
        remove_paragraph(p)


def get_next_subfolder(base: Path) -> Path:
    """Find the next numeric subfolder under base, e.g., 1, 2, 3..."""
    base.mkdir(parents=True, exist_ok=True)
    numbers = []
    for child in base.iterdir():
        if child.is_dir() and re.fullmatch(r"\d+", child.name):
            numbers.append(int(child.name))
    next_num = max(numbers, default=0) + 1
    return base / str(next_num)


def main():
    """Main resume generation logic."""
    if not TEMPLATE_PATH.exists():
        print(f"Template not found: {TEMPLATE_PATH.resolve()}")
        sys.exit(1)
    if not JSON_PATH.exists():
        print(f"JSON not found: {JSON_PATH.resolve()}")
        sys.exit(1)

    with JSON_PATH.open("r", encoding="utf-8") as f:
        ctx = json.load(f)

    # --------------------------
    # Prepare numbered folder
    # --------------------------
    target_folder = get_next_subfolder(BASE_FOLDER)
    target_folder.mkdir(parents=True, exist_ok=True)

    # Define final paths inside that folder
    OUTPUT_RAW = target_folder / "Tailored_Resume_RAW.docx"
    OUTPUT_FINAL = target_folder / FINAL_FILENAME

    # --------------------------
    # Render raw (template render)
    # --------------------------
    doc = DocxTemplate(str(TEMPLATE_PATH))
    doc.render(ctx)
    doc.save(str(OUTPUT_RAW))

    # --------------------------
    # Post-process for final (cleanup + formatting)
    # --------------------------
    d = Document(str(OUTPUT_RAW))
    normalize_bullets(d)

    for p in d.paragraphs:
        if is_list_style(p):
            tighten_list_paragraph(p)

    d.save(str(OUTPUT_FINAL))

    print(f"[OK] Rendered OK: {OUTPUT_FINAL.resolve()}")
    print(f"  (Raw pre-fix version at: {OUTPUT_RAW.resolve()})")


if __name__ == "__main__":
    main()
